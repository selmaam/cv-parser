
    
import streamlit as st
import pandas as pd
import os
import sqlite3
import translation
import text_extraction
import extracting_information
import preprocessing
import main_test
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import visualization
import json
import matching
import resume_information
import subprocess
import spacy 
import update_skills
import matchers
import time

nlp = spacy.load("en_core_web_sm")

# Authentication code
users = {
    "admin": "password123"
}

# Function to authenticate user
def authenticate_user(username, password):
    return users.get(username) == password

def login_page():
    st.markdown(
        """
        <style>
        .login-container {
            max-width: 600px;
            margin: auto;
            padding: 2rem;
            border-radius: 10px;
            background-color: #FF0000;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
            text-align: center;
            display: flex;
            justify-content: center;
            align-items: center;
            flex-wrap: wrap;
        }
        .login-title {
            font-size: 28px;
            margin-bottom: 1.5rem;
            color: #FF0000;
            text-align: center;
            width: 100%;
        }
        .login-form {
            width: 50%;
            display: flex;
            justify-content: center;
            align-items: center;
            flex-wrap: wrap;
        }
        .login-input {
            margin-bottom: 1rem;
            padding: 0.5rem;
            border: 1px solid #cccccc;
            border-radius: 5px;
            flex: 1;
            max-width: 100px;
            height: 40px; /* Add this to set the height of the input fields */
        }

        .login-button:hover {
            background-color: #357abd;
        }
        .error-message {
            color: #ff6347;
            text-align: center;
            margin-top: 1rem;
            width: 100%;
        }
        </style>
        """,
        unsafe_allow_html=True)

    # st.markdown('<div class="login-container">', unsafe_allow_html=True)
    st.markdown('<div class="login-title">Login</div>', unsafe_allow_html=True)

    with st.form("login_form"):
        username = st.text_input('Username', key='username_input', placeholder='Enter your username', max_chars=30, help='Enter your username')
        password = st.text_input('Password', type='password', key='password_input', placeholder='Enter your password', max_chars=30, help='Enter your password')
        submit_button = st.form_submit_button(label='Login')

    if submit_button:
        if authenticate_user(username, password):
            st.session_state.logged_in = True  # Set session state for redirection
            st.session_state.username = username
            st.experimental_rerun()
        else:
            st.markdown('<div class="error-message">Invalid username or password</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)
    

    
def update_database(database_file, max_chars_per_column=50):
    # Connect to the SQLite database
    conn = sqlite3.connect(database_file)
    cursor = conn.cursor()

    # Get the list of tables in the database
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()

    st.title("Database Viewer")

    if 'show_add_form' not in st.session_state:
        st.session_state.show_add_form = {table[0]: False for table in tables}
    if 'show_delete_form' not in st.session_state:
        st.session_state.show_delete_form = {table[0]: False for table in tables}

    for table in tables:
        table_name = table[0]
        cursor.execute(f"SELECT * FROM {table_name};")
        table_contents = cursor.fetchall()
        column_names = [description[0] for description in cursor.description]

        # Create a DataFrame from the table contents
        df = pd.DataFrame(table_contents, columns=column_names)

        # Truncate text in each cell if it exceeds max_chars_per_column
        truncated_df = df.applymap(lambda x: truncate_text(str(x), max_chars_per_column))

        # Select the column 'file_name' to display
        file_name_column = 'file_name'
        if file_name_column in truncated_df.columns:
            st.dataframe(truncated_df[file_name_column], width=400,)

        # Create columns for buttons
        col1, col2 = st.columns(2)

        # Button to show the add form
        if col1.button(f"Add to {table_name}", key=f"add_{table_name}"):
            st.session_state.show_add_form[table_name] = not st.session_state.show_add_form[table_name]

        # Button to show the delete form
        if col2.button(f"Delete from {table_name}", key=f"delete_{table_name}"):
            st.session_state.show_delete_form[table_name] = not st.session_state.show_delete_form[table_name]

        # Show the add form if the button was clicked
        if st.session_state.show_add_form[table_name]:
            with st.form(key=f"add_form_{table_name}"):
                st.subheader(f"Add a new entry to {table_name}")
                uploaded_files = st.file_uploader("Choose files", type=["pdf", "docx", "txt", "jpg", "jpeg", "png"], accept_multiple_files=True, key=f"upload_{table_name}")
                submit_button = st.form_submit_button(label=f"Add to {table_name}")
                if submit_button and uploaded_files:
                    for uploaded_file in uploaded_files:
                        # Add the uploaded file to the database
                        add_files_to_db(uploaded_file)
                    st.success(f"New entries added to {table_name}!")
                    st.session_state.show_add_form[table_name] = False

        # Show the delete form if the button was clicked
        if st.session_state.show_delete_form[table_name]:
            with st.form(key=f"delete_form_{table_name}"):
                file_names = [row[1] for row in table_contents]  # Assuming the second column is file names
                files_to_delete = st.multiselect(f"Select file names to delete from {table_name}:", file_names, key=f"delete_files_{table_name}")
                delete_button = st.form_submit_button(label=f"Delete from {table_name}")
                if delete_button and files_to_delete:
                    for file_to_delete in files_to_delete:
                        # Construct the DELETE query based on the selected file name
                        query = f"DELETE FROM {table_name} WHERE file_name=?"
                        cursor.execute(query, (file_to_delete,))
                    conn.commit()
                    st.success(f"Rows with selected file names deleted successfully from {table_name}!")
                    st.session_state.show_delete_form[table_name] = False

                    # Refresh the displayed table
                    cursor.execute(f"SELECT * FROM {table_name};")
                    updated_table_contents = cursor.fetchall()
                    updated_df = pd.DataFrame(updated_table_contents, columns=column_names)
                    truncated_updated_df = updated_df.applymap(lambda x: truncate_text(str(x), max_chars_per_column))
                    st.dataframe(truncated_updated_df, width=800, height=400)

        st.write("---")  # Add a separator between each table

    # Close the database connection
    conn.close()   
    
    


def details(text, text1):
    
    information = resume_information.extract_essential_infos(text, text1)
    return information


def view_tables_and_contents(database_file, skills, max_chars_per_column=50):
    
    # Connect to the SQLite database
    conn = sqlite3.connect(database_file)
    cursor = conn.cursor()
    st.session_state.main_action = "View Resume Database"
    # Get the list of tables in the database
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()
    

    st.title("Database Viewer")

    for table in tables:
        table_name = table[0]
        cursor.execute(f"SELECT * FROM {table_name};")
        table_contents = cursor.fetchall()
        column_names = [description[0] for description in cursor.description]

        # Create a DataFrame from the table contents
        df = pd.DataFrame(table_contents, columns=column_names)

        # Truncate text in each cell if it exceeds max_chars_per_column
        truncated_df = df.applymap(lambda x: truncate_text(str(x), max_chars_per_column))

        # Select the column 'file_name' to display
        file_name_column = 'file_name'
        if file_name_column in truncated_df.columns:
            st.dataframe(truncated_df[file_name_column], width=400,)
            
        for table in tables:
            table_name = table[0]
            query = f"SELECT * FROM {table_name};"
            df = pd.read_sql_query(query, conn)
            columns_to_drop = ['file_path', 'id']  
            df.drop(columns=columns_to_drop, inplace=True)
            processed_data = df["processed"]
            skills_list = []
            score_cv_scores = []
            score_cv_skills = []


       
        st.session_state.option1 = 'Choose option'
        st.session_state.option = 'Choose option'
        selected_option1 = 'Choose option'
        # print(st.session_state.option)
        
        st.title("Choose Display Option")
        if 'option' not in st.session_state:
            st.session_state.option1 = 'Choose option'
        
        
        # print(st.session_state.option1)
        # Store the selected option in session state
        selected_option1 = st.selectbox('Display Options:', ['Choose option', 'Threshold', 'Top Matches'], index=['Choose option', 'Threshold', 'Top Matches'].index(st.session_state.option1))
        # print(selected_option1)
        if selected_option1 != 'Choose option' and selected_option1 != st.session_state.option1:
            st.session_state.option1 = selected_option1
        # print(selected_option1)
        # print(st.session_state.option1)                                  
        if st.session_state.option1 == 'Top Matches' :
                    for processed_item in processed_data:
                        cvs_data = dataa[dataa['full'].isin(skills['full'])]
                        skills_df = matchers.skills_extraction_pipeline(cvs_data, processed_item, full_matcher, abbv_matcher, nlp)
                        # df_match = skills_df[["full", "type", "score"]]
                        

                        # Create a dictionary for the skills
                        skills_dict = {row["full"]: (row["type"], row["score"]) for _, row in skills_df.iterrows()}
                        skills_list.append(skills_dict)
                                                
                        match_score = matching.check_cv_match(skills, skills_df)['match_percentage']
                        score_cv_scores.append(match_score if match_score is not None else 0)
                        # match_skills = matching.check_cv_match(skills, skills_df)['matched_skills']
                        # score_cv_skills.append(match_skills)
                        # print(match_skills)
                    
                    
                    # Add the skills_list as a new column named "skills" to the DataFrame
                    df['skills'] = skills_list
                    df['match_score'] = score_cv_scores
                    # df['common_skills'] = score_cv_skills
                    
                    # Sort the DataFrame by 'match_score' column in descending order
                    df_sorted = df.sort_values(by='match_score', ascending=False)
                    # print(df_sorted)
                    
                    top_5_df = df_sorted.head(5)
                    # print(top_5_df)
                    st.header(f"Top matches ")

                    # Make file names clickable for details
                    for index, row in top_5_df.iterrows():

                        with st.expander(f"Details for {row['file_name']}"):
                            st.markdown("""
                            <style>
                            .scrollable-container {
                                max-height: 400px; /* Adjust the height as needed */
                                overflow-y: scroll;
                                padding-right: 15px; /* Add padding to avoid cutting off content */
                            }
                            .scrollable-container::-webkit-scrollbar {
                                width: 8px;
                            }
                            .scrollable-container::-webkit-scrollbar-thumb {
                                background-color: #888;
                                border-radius: 10px;
                            }
                            .scrollable-container::-webkit-scrollbar-thumb:hover {
                                background-color: #555;
                            }
                            </style>
                            """, unsafe_allow_html=True)

                            # Wrap the content in a scrollable container
                            st.markdown('<div class="scrollable-container">', unsafe_allow_html=True)
                            file_name = row['file_name']
                            upload_dir = 'uploaded_DB'
                            upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                            # st.markdown(upload_dir_path)
                            file_path = os.path.join(upload_dir_path, file_name)     
                            # st.markdown(file_path)                                           
                            # st.markdown(file_path)
                            if os.path.exists(file_path):
                                with open(file_path, "rb") as file:
                                    file_data = file.read()
                                st.download_button(label="Open PDF", data=file_data, file_name=file_name, mime="application/pdf")
                            else:
                                st.write("File not found.")
                            
                            st.write(f"**Match Score:** {row['match_score']}")
                            
                            st.markdown('</div>', unsafe_allow_html=True)

                            processed_text = row['processed']
                            essential_info = details(processed_text, row['file_name'])

                            st.markdown("""
                            <style>
                            .dashboard-container {
                                padding: 10px;
                                background-color: #f0f0f0;
                                border-radius: 5px;
                                box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                margin-top: 20px; /* Add margin to separate from the scrollable container */
                            }
                            .dashboard-title {
                                font-size: 20px;
                                font-weight: bold;
                                margin-bottom: 10px;
                            }
                            .dashboard-item {
                                margin-bottom: 5px;
                            }
                            </style>
                            """, unsafe_allow_html=True)

                            st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
                            st.markdown('<div class="dashboard-title">Essential Information</div>', unsafe_allow_html=True)
                            for key, value in essential_info.items():
                                if key == 'education' and value:
                                    st.markdown('<div class="dashboard-item"><b>Education:</b></div>', unsafe_allow_html=True)
                                    education_df = pd.DataFrame(value, columns=['Education'])
                                    st.dataframe(education_df, width=800)
                                elif value:
                                    st.markdown(f'<div class="dashboard-item"><b>{key}:</b> {value}</div>', unsafe_allow_html=True)
                                else:
                                    st.markdown(f'<div class="dashboard-item"><b>{key}:</b> Not available</div>', unsafe_allow_html=True)
                            st.markdown('</div>', unsafe_allow_html=True)
                            
                            skills_dict = row['skills']
                            skills_df = pd.DataFrame.from_dict(skills_dict, orient='index', columns=['type', 'score'])
                            # Create a DataFrame with skills from the dictionary
                            skills_df = skills_df.sort_values(by='score', ascending=False)
                            skills_df_only = pd.DataFrame({'skills': list(skills_dict.keys())})

                            # Drop the 'type' column
                            skills_df = skills_df.drop(columns=['type'])  # This line drops the 'type' column

                            st.subheader("Common Skills with job descritption")
                            st.dataframe(skills_df, width=800)
                            # Display matched skills DataFrame
                            # Ensure matched_skills_df is defined and populated
                            # if 'common_skills' in row and row['common_skills'] is not None:
                            #     common_skills_list = row['common_skills']
                            #     common_skills_df = pd.DataFrame([common_skills_list], columns=[f"Skill {i+1}" for i in range(len(common_skills_list))])
                                # st.subheader("Common Skills with job descritption")
                                # st.dataframe(common_skills_df)

                            st.markdown("""
                            <style>
                            .dashboard-container {
                                padding: 10px;
                                background-color: #ff9999; /* Light red shade */
                                border-radius: 5px;
                                box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                margin-top: 20px; /* Add margin to separate from the scrollable container */
                            }
                            .dashboard-title {
                                font-size: 20px;
                                font-weight: bold;
                                margin-bottom: 10px;
                            }
                            .dashboard-item {
                                margin-bottom: 5px;
                            }
                            </style>
                            """, unsafe_allow_html=True)

                            # Move the pie chart below the essential information and align it to the left
                            # skills_from_job_desc = pd.DataFrame(skills['full'], columns=['skills'])

            # Move the pie chart below the essential information and align it to the left
                            fig_pie = visualization.display_skills_distribution(skills_df_only, cvs_data)
                            st.markdown('<div style="clear:both;"></div>', unsafe_allow_html=True)  # Clear float
                            st.plotly_chart(fig_pie, use_container_width=False)
                            st.session_state.main_action = None
                            st.session_state.option = 'Choose option'
                                
        elif st.session_state.option1 == 'Threshold' :
            # threshold_input = st.number_input("Enter the threshold score:", value=0, step=0.01)
                for processed_item in processed_data:
                    cvs_data = dataa[dataa['full'].isin(skills['full'])]
                    skills_df = matchers.skills_extraction_pipeline(cvs_data, processed_item, full_matcher, abbv_matcher, nlp)
                    # df_match = skills_df[["full", "type", "score"]]
                    

                    # Create a dictionary for the skills
                    skills_dict = {row["full"]: (row["type"], row["score"]) for _, row in skills_df.iterrows()}
                    skills_list.append(skills_dict)
                                            
                    match_score = matching.check_cv_match(skills, skills_df)['match_percentage']
                    score_cv_scores.append(match_score if match_score is not None else 0)
                    match_skills = matching.check_cv_match(skills, skills_df)['matched_skills']
                    score_cv_skills.append(match_skills)
                    # print(match_skills)
                
                
                # Add the skills_list as a new column named "skills" to the DataFrame
                df['skills'] = skills_list
                df['match_score'] = score_cv_scores
                df['common_skills'] = score_cv_skills
                
                # Sort the DataFrame by 'match_score' column in descending order
                df_sorted = df.sort_values(by='match_score', ascending=False)
                print(df_sorted)
                threshold_df = df_sorted[df_sorted['match_score'] >= 0.2]
                if threshold_df.empty:
                    st.write("No matches above the threshold were found.")
                else:
                    
                        with st.expander("View Match Statistics"):
                            number_of_resumes = len(df_sorted)
                            number_of_matches = len(threshold_df)
                            
                            fig_pie = visualization.plot_matches_pie(number_of_resumes, number_of_matches)

                            # Display the pie chart in Streamlit
                            st.plotly_chart(fig_pie)
                            
                        st.header(f"Matches above Threshold ")    
                        threshold_df = threshold_df.sort_values(by='match_score', ascending=False)                                        
                        for index, row in threshold_df.iterrows():

                        
                           
                            with st.expander(f"Details for {row['file_name']}"):
                                    st.markdown("""
                                    <style>
                                    .scrollable-container {
                                        max-height: 400px; /* Adjust the height as needed */
                                        overflow-y: scroll;
                                        padding-right: 15px; /* Add padding to avoid cutting off content */
                                    }
                                    .scrollable-container::-webkit-scrollbar {
                                        width: 8px;
                                    }
                                    .scrollable-container::-webkit-scrollbar-thumb {
                                        background-color: #888;
                                        border-radius: 10px;
                                    }
                                    .scrollable-container::-webkit-scrollbar-thumb:hover {
                                        background-color: #555;
                                    }
                                    </style>
                                    """, unsafe_allow_html=True)

                                    # Wrap the content in a scrollable container
                                    st.markdown('<div class="scrollable-container">', unsafe_allow_html=True)
                                    file_name = row['file_name']
                                    upload_dir = 'uploaded_DB'
                                    upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                                    # st.markdown(upload_dir_path)
                                    file_path = os.path.join(upload_dir_path, file_name)     
                                    # st.markdown(file_path)                                           
                                    # st.markdown(file_path)
                                    if os.path.exists(file_path):
                                        with open(file_path, "rb") as file:
                                            file_data = file.read()
                                        st.download_button(label="Open PDF", data=file_data, file_name=file_name, mime="application/pdf")
                                    else:
                                        st.write("File not found.")
                                    st.write(f"**Match Score:** {row['match_score']}")
                                    
                                    st.markdown('</div>', unsafe_allow_html=True)

                                    processed_text = row['processed']
                                    essential_info = details(processed_text, row['file_name'])

                                    st.markdown("""
                                    <style>
                                    .dashboard-container {
                                        padding: 10px;
                                        background-color: #f0f0f0;
                                        border-radius: 5px;
                                        box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                        margin-top: 20px; /* Add margin to separate from the scrollable container */
                                    }
                                    .dashboard-title {
                                        font-size: 20px;
                                        font-weight: bold;
                                        margin-bottom: 10px;
                                    }
                                    .dashboard-item {
                                        margin-bottom: 5px;
                                    }
                                    </style>
                                    """, unsafe_allow_html=True)

                                    st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
                                    st.markdown('<div class="dashboard-title">Essential Information</div>', unsafe_allow_html=True)
                                    for key, value in essential_info.items():
                                        if key == 'education' and value:
                                            st.markdown('<div class="dashboard-item"><b>Education:</b></div>', unsafe_allow_html=True)
                                            education_df = pd.DataFrame(value, columns=['Education'])
                                            st.dataframe(education_df, width=800)
                                        elif value:
                                            st.markdown(f'<div class="dashboard-item"><b>{key}:</b> {value}</div>', unsafe_allow_html=True)
                                        else:
                                            st.markdown(f'<div class="dashboard-item"><b>{key}:</b> Not available</div>', unsafe_allow_html=True)
                                    st.markdown('</div>', unsafe_allow_html=True)
                                    
                                    skills_dict = row['skills']
                                    skills_df = pd.DataFrame.from_dict(skills_dict, orient='index', columns=['type', 'score'])
                                    # Create a DataFrame with skills from the dictionary
                                    skills_df = skills_df.sort_values(by='score', ascending=False)
                                    skills_df_only = pd.DataFrame({'skills': list(skills_dict.keys())})

                                    # Drop the 'type' column
                                    skills_df = skills_df.drop(columns=['type'])  # This line drops the 'type' column

                                    st.subheader("Common Skills with job descritption")
                                    st.dataframe(skills_df, width=800)
                                    # Display matched skills DataFrame
                                    # # Ensure matched_skills_df is defined and populated
                                    # if 'common_skills' in row and row['common_skills'] is not None:
                                    #     common_skills_list = row['common_skills']
                                    #     common_skills_df = pd.DataFrame([common_skills_list], columns=[f"Skill {i+1}" for i in range(len(common_skills_list))])
                                        # st.subheader("Common Skills with job descritption")
                                        # st.dataframe(common_skills_df)

                                    st.markdown("""
                                    <style>
                                    .dashboard-container {
                                        padding: 10px;
                                        background-color: #ff9999; /* Light red shade */
                                        border-radius: 5px;
                                        box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                        margin-top: 20px; /* Add margin to separate from the scrollable container */
                                    }
                                    .dashboard-title {
                                        font-size: 20px;
                                        font-weight: bold;
                                        margin-bottom: 10px;
                                    }
                                    .dashboard-item {
                                        margin-bottom: 5px;
                                    }
                                    </style>
                                    """, unsafe_allow_html=True)
                                    

            # Move the pie chart below the essential information and align it to the left
                                    fig_pie = visualization.display_skills_distribution(skills_df_only, cvs_data)
                                    # Move the pie chart below the essential information and align it to the left
                                    # fig_pie = visualization.display_skills_distribution(skills_df_only, skills['full'])
                                    st.markdown('<div style="clear:both;"></div>', unsafe_allow_html=True)  # Clear float
                                    st.plotly_chart(fig_pie, use_container_width=False)
                                    st.session_state.main_action = None
                                    st.session_state.option = 'Choose option'
                                    

            # Close the connection
                conn.close()




def process_uploaded_pdfs(uploaded_files):
    """Process uploaded files and return a DataFrame with extracted text."""
    temp_dir = "temp"
    os.makedirs(temp_dir, exist_ok=True)
    data = []
    
    for uploaded_file in uploaded_files:
        # Save the uploaded file to a temporary directory
        temp_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Extract text based on the file type
        if temp_path.endswith(".pdf"):
            # Extract text from PDF using OCR
            # text = OCR_extraction.extract_text_from_resume(temp_path)
            text = text_extraction.extract_text_from_pdf(temp_path)
        elif temp_path.endswith(".docx"):
            # Extract text from Word document
            doc = Document(temp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif temp_path.endswith(".txt"):
            # Read text from a plain text file
            with open(temp_path, "r") as txt_file:
                text = txt_file.read()
        elif temp_path.endswith((".jpg", ".jpeg", ".png")):
            # Extract text from image using OCR (Pytesseract)
            image = Image.open(temp_path)
            # text = text_extraction.extract_text_from_pdf(image)
            text = pytesseract.image_to_string(image)
        else:
            text = "Unsupported file type"
        
        # Perform any additional processing, such as translation
        text = translation.translation(text)
        
        # Add file information and extracted text to the data list
        data.append({"file_name": uploaded_file.name, "text": text})
        
        # Remove the temporary file after processing
        os.remove(temp_path)
    
    # Create a DataFrame from the extracted data
    df = pd.DataFrame(data)
    
    return df


full_matcher, abbv_matcher, dataa, nlp = matchers.load_matchers_and_data()

def load_tree():
    file_path = "final_raw_skills.json"
    try:
        with open(file_path, 'r') as file:
            data = json.load(file)
    except FileNotFoundError:
        st.error(f"Error: File '{file_path}' not found.")
        data = None
    except Exception as e:
        st.error(f"Error reading file: {e}")
        data = None
    return data

def run_script(nlp):
    update_skills.Fetch_EMSISkills()
    update_skills.hierarchical_data(input_file="raw_skills.json", output_file="updated_skills.json")
    update_skills.skills_preparation(file_path='updated_skills.json')
    update_skills.matchers_creation(nlp, pd.read_csv('skills.csv'))
    #subprocess.run(["python", "update_skills.py"])


def show_main_interface():
    st.title("Job Description and CV Processor")
    # st.session_state.main_action = None
    
    if 'nav_choice' not in st.session_state:
        st.session_state.nav_choice = None
    if 'main_action' not in st.session_state:
        st.session_state.main_action = None
        
    if 'main_action_form' not in st.session_state:
        st.session_state.main_action_form = None
    

    data = load_tree() 

    st.markdown("""
    <style>
    div.stButton > button {
        width: 100%;
        margin: auto;
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <style>
    div.stSidebar {
        padding: 20px;
        border: 1px solid #ddd;
        border-radius: 10px;
        box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
    }

    div.stSidebar > div:nth-child(1) {
        font-size: 18px;
        font-weight: bold;
        margin-bottom: 10px;
        border-bottom: 1px solid #ddd;
        padding-bottom: 10px;
    }

    div.stSidebar > div:not(:first-child) {
        margin-top: 20px;
    }

    div.stButton > button {
        padding: 10px 20px;
        font-size: 16px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
    }

    div.stButton > button:hover {
        background-color: #f7f7f7;
    }
    </style>
    """, unsafe_allow_html=True)

    # Sidebar navigation
    st.sidebar.title("Navigation Bar")

    if st.sidebar.button("Upload Job Description"):
        st.session_state.nav_choice = "Upload Job Description"
        
    if st.sidebar.button("Fill Job Description Form"):
        st.session_state.nav_choice = "Fill Job Description Form"

    if st.sidebar.button("Update Resume Database"):
        st.session_state.nav_choice = "Update Resume Database"

    if st.sidebar.button("Additional actions"):
        st.session_state.nav_choice = "Additional actions"
        
    # st.header("Upload Job Description")
    
    # Display introductory content or features overview
    if st.session_state.nav_choice is None: #or st.session_state.nav_choice == "Additional actions":
     st.markdown("""
    ## Welcome to the Job Description and CV Processor App!

    This app helps you streamline the process of analyzing job descriptions and CVs.

    **Key Features:**
    - Upload job descriptions for skill extraction.
    - Fill out a form to specify job requirements.
    - extract skills from job description and resumes.
    - rank resumes based on skills 
    - extract personal information 
    - Perform additional actions like updating skills database.

    **Get Started:**
    Choose an option from the sidebar to begin.
    """)

    # Determine the current navigation choice
    nav_choice = st.session_state.nav_choice
    
    
     # Clear the content of the screen
    upload_placeholder = st.empty()
    # Placeholder for Fill Job Description Form section
    form_placeholder = st.empty()
    
    if nav_choice == "Upload Job Description":
        
                upload_placeholder.empty()
                form_placeholder.empty()
        
                st.header("Upload Job Description")
                job_desc_file = st.file_uploader("Upload a job description file", type=["txt"])

                if job_desc_file:
                #  if 'extracted_skills' not in st.session_state or 'extracted_res' not in st.session_state:
                    job_desc_bytes = job_desc_file.read()
                    job_desc = job_desc_bytes.decode("utf-8")
                    job_desc = translation.translation(job_desc)
                    # processed_desc = main_test.processing_job(job_desc)
                    
                    st.session_state.extracted_skills = main_test.get_skills(job_desc)  # Extract skills
                    # print("SKILLLLLS")
                   
                    st.session_state.extracted_skills = st.session_state.extracted_skills.drop_duplicates(subset='full')
                    # print(st.session_state.extracted_skills)
                    # st.session_state.skills = st.session_state.extracted_skills.drop(columns=['ngram'])
                    sentences = preprocessing.tokenize_sentences(preprocessing.special_char(job_desc))
                    # st.write(sentences)
                    st.session_state.extracted_res = extracting_information.extract_resp(sentences)  # Extract responsibilities
                    

                    # Display file name and main interface buttons
                    st.markdown(f"**File Name:** {job_desc_file.name}")
                    col1, col2, col3 = st.columns(3)
                    # st.session_state.main_action  = None
                    if col1.button("View Resume Database"):
                        st.session_state.main_action = "View Resume Database"
                    if col2.button("Upload Resumes"):
                        st.session_state.main_action = "Upload Resumes"
                    if col3.button("Details"):
                        st.session_state.main_action = "Details"

                    main_action = st.session_state.main_action
                    st.session_state.submitted = False
                    st.session_state.form_data = None
                    st.session_state.main_action2 = None
                            

                    # Execute main actions based on button clicks
                    if main_action == "View Resume Database":
                        st.session_state.option = 'Choose option'
                        
                        database_file = "pdfs.db"  # Replace this with your actual database file path
                        view_tables_and_contents(database_file, st.session_state.extracted_skills)
                        
                        
                        
                        
                    if main_action == "Upload Resumes":
                        st.header("Upload Resumes")
                        st.session_state.option1 = 'Choose option'
                        selected_option1 = 'Choose option'
                        uploaded_files = st.file_uploader("Choose a file", accept_multiple_files=True, type=["pdf", "docx", "txt", "jpg", "jpeg", "png"])
                        
                        
                        
                        if uploaded_files:
                            
                            upload_dir = 'uploaded_files'
                            if not os.path.exists(upload_dir):
                                os.makedirs(upload_dir)
                            
                            for uploaded_file in uploaded_files:
                                upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                                                        # st.markdown(upload_dir_path)
                                file_path = os.path.join(upload_dir_path, uploaded_file.name)
                                # file_path = os.path.join(upload_dir, uploaded_file.name)
                                with open(file_path, 'wb') as file:
                                    file.write(uploaded_file.read())
                                
                            start_time = time.time()
                            df = process_uploaded_pdfs(uploaded_files)
                            # updated_df = view_resume_data(df)
                            st.write(df)
                        
                        
                            

                            # Display the updated DataFrame
                            st.success(f"Processing completed!")
                            st.session_state.option = 'Choose option'
                            st.title("Choose Display Option")
                            if 'option' not in st.session_state:
                                st.session_state.option = 'Choose option'

                            # Store the selected option in session state
                            selected_option = st.selectbox('Display Options:', ['Choose option', 'Threshold', 'Top Matches'], index=['Choose option', 'Threshold', 'Top Matches'].index(st.session_state.option))
                            if selected_option != 'Choose option':
                                st.session_state.option = selected_option
                                                                
                            if st.session_state.option == 'Top Matches':
                                        processed_data = df["text"]
                                        skills_list = []
                                        score_cv_scores = []
                                        score_cv_skills = []

                                        # Process each item using list comprehension
                                        # processed_result = [main_test.processing_job(item) for item in processed_data]
                                        
                                        cvs_data = dataa[dataa['full'].isin(st.session_state.extracted_skills['full'])]
                                        # processed_result = [main_test.processing_job(item) for item in processed_data]

                                        # Extract skills from each processed item and add to skills_list
                                        for processed_item in processed_data:
                                            skills_df = matchers.skills_extraction_pipeline(cvs_data, processed_item, full_matcher, abbv_matcher, nlp)
                                            # df_match = skills_df[["full", "type", "score"]]

                                            # Create a dictionary for the skills
                                            skills_dict = {row["full"]: (row["type"], row["score"]) for _, row in skills_df.iterrows()}
                                            skills_list.append(skills_dict)

                                            match_score = matching.check_cv_match(st.session_state.extracted_skills, skills_df)['match_percentage']
                                            # print("match score", match_score)
                                            score_cv_scores.append(match_score if match_score is not None else 0)
                                            match_skills = matching.check_cv_match(st.session_state.extracted_skills, skills_df)['matched_skills']
                                            score_cv_skills.append(match_skills)

                                        # Add the skills_list as a new column named "skills" to the DataFrame
                                        df['skills'] = skills_list
                                        df['match_score'] = score_cv_scores
                                        df['common_skills'] = score_cv_skills
                                        
                                        # st.session_state.option = 'Choose option'

                                        # Sort the DataFrame by 'match_score' column in descending order
                                        df_sorted = df.sort_values(by='match_score', ascending=False)

                                        top_5_df = df_sorted.head(5)
                                        st.header(f"Top matches ")

                                        # Make file names clickable for details
                                        for index, row in top_5_df.iterrows():
                                            # st.markdown(f"[Open PDF]({row['file_path']})")
                                            

                                            with st.expander(f"Details for {row['file_name']}"):
                                                st.markdown("""
                                                <style>
                                                .scrollable-container {
                                                    max-height: 400px; /* Adjust the height as needed */
                                                    overflow-y: scroll;
                                                    padding-right: 15px; /* Add padding to avoid cutting off content */
                                                }
                                                .scrollable-container::-webkit-scrollbar {
                                                    width: 8px;
                                                }
                                                .scrollable-container::-webkit-scrollbar-thumb {
                                                    background-color: #888;
                                                    border-radius: 10px;
                                                }
                                                .scrollable-container::-webkit-scrollbar-thumb:hover {
                                                    background-color: #555;
                                                }
                                                </style>
                                                """, unsafe_allow_html=True)

                                                # Wrap the content in a scrollable container
                                                st.markdown('<div class="scrollable-container">', unsafe_allow_html=True)
                                                st.markdown(f"**Match Score:** {row['match_score']}")
                                                # Make the file name clickable to open the PDF file
                                                file_name = row['file_name']
                                                upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                                                # st.markdown(upload_dir_path)
                                                file_path = os.path.join(upload_dir_path, file_name)     
                                                # st.markdown(file_path)                                           
                                                # st.markdown(file_path)
                                                if os.path.exists(file_path):
                                                    with open(file_path, "rb") as file:
                                                        file_data = file.read()
                                                    st.download_button(label="Open PDF", data=file_data, file_name=file_name, mime="application/pdf")
                                                else:
                                                    st.write("File not found.")
                                                
                                                st.markdown('</div>', unsafe_allow_html=True)

                                                processed_text = row['text']
                                                essential_info = details(processed_text, row['file_name'])

                                                st.markdown("""
                                                <style>
                                                .dashboard-container {
                                                    padding: 10px;
                                                    background-color: #f0f0f0;
                                                    border-radius: 5px;
                                                    box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                                    margin-top: 20px; /* Add margin to separate from the scrollable container */
                                                }
                                                .dashboard-title {
                                                    font-size: 20px;
                                                    font-weight: bold;
                                                    margin-bottom: 10px;
                                                }
                                                .dashboard-item {
                                                    margin-bottom: 5px;
                                                }
                                                </style>
                                                """, unsafe_allow_html=True)

                                                st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
                                                st.markdown('<div class="dashboard-title">Essential Information</div>', unsafe_allow_html=True)
                                                for key, value in essential_info.items():
                                                    if key == 'education' and value:
                                                        st.markdown('<div class="dashboard-item"><b>Education:</b></div>', unsafe_allow_html=True)
                                                        education_df = pd.DataFrame(value, columns=['Education'])
                                                        st.dataframe(education_df, width=800)
                                                    elif value:
                                                        st.markdown(f'<div class="dashboard-item"><b>{key}:</b> {value}</div>', unsafe_allow_html=True)
                                                    else:
                                                        st.markdown(f'<div class="dashboard-item"><b>{key}:</b> Not available</div>', unsafe_allow_html=True)
                                                st.markdown('</div>', unsafe_allow_html=True)
                                                
                                                skills_dict = row['skills']
                                                skills_df = pd.DataFrame.from_dict(skills_dict, orient='index', columns=['type', 'score'])
                                                # Create a DataFrame with skills from the dictionary
                                                skills_df = skills_df.sort_values(by='score', ascending=False)
                                                skills_df_only = pd.DataFrame({'skills': list(skills_dict.keys())})

                                                # Drop the 'type' column
                                                skills_df = skills_df.drop(columns=['type'])  # This line drops the 'type' column

                                                st.subheader("Common Skills with job descritption")
                                                st.dataframe(skills_df, width=800)
                                                # Display matched skills DataFrame
                                                # Ensure matched_skills_df is defined and populated
                                                if 'common_skills' in row and row['common_skills'] is not None:
                                                    common_skills_list = row['common_skills']
                                                    common_skills_df = pd.DataFrame([common_skills_list], columns=[f"Skill {i+1}" for i in range(len(common_skills_list))])
                                                    # st.subheader("Common Skills with job descritption")
                                                    # st.dataframe(common_skills_df)

                                                st.markdown("""
                                                <style>
                                                .dashboard-container {
                                                    padding: 10px;
                                                    background-color: #ff9999; /* Light red shade */
                                                    border-radius: 5px;
                                                    box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                                    margin-top: 20px; /* Add margin to separate from the scrollable container */
                                                }
                                                .dashboard-title {
                                                    font-size: 20px;
                                                    font-weight: bold;
                                                    margin-bottom: 10px;
                                                }
                                                .dashboard-item {
                                                    margin-bottom: 5px;
                                                }
                                                </style>
                                                """, unsafe_allow_html=True)

                                                # Move the pie chart below the essential information and align it to the left
                                                fig_pie = visualization.display_skills_distribution(skills_df_only, cvs_data)
                                                st.markdown('<div style="clear:both;"></div>', unsafe_allow_html=True)  # Clear float
                                                st.plotly_chart(fig_pie, use_container_width=False)
                                                st.session_state.main_action = None
                                                st.session_state.option = 'Choose option'
                                                st.session_state.option1 = 'Choose option'
                                                  
                            elif st.session_state.option == 'Threshold':
                                # threshold_input = st.number_input("Enter the threshold score:", value=0, step=0.01)

                                # IM HEEEEEEEEEEEEEEERE 
                                    processed_data = df["text"]
                                    skills_list = []
                                    score_cv_scores = []
                                    score_cv_skills = []

                                    # Process each item using list comprehension
                                    # processed_result = [main_test.processing_job(item) for item in processed_data]
                                    
                                    cvs_data = dataa[dataa['full'].isin(st.session_state.extracted_skills['full'])]
                                    # processed_result = [main_test.processing_job(item) for item in processed_data]

                                    # Extract skills from each processed item and add to skills_list
                                    for processed_item in processed_data:
                                        skills_df = matchers.skills_extraction_pipeline(cvs_data, processed_item, full_matcher, abbv_matcher, nlp)
                                        # df_match = skills_df[["full", "type", "score"]]

                                        # Create a dictionary for the skills
                                        skills_dict = {row["full"]: (row["type"], row["score"]) for _, row in skills_df.iterrows()}
                                        skills_list.append(skills_dict)

                                        match_score = matching.check_cv_match(st.session_state.extracted_skills, skills_df)['match_percentage']
                                        # print("match score", match_score)
                                        score_cv_scores.append(match_score if match_score is not None else 0)
                                        match_skills = matching.check_cv_match(st.session_state.extracted_skills, skills_df)['matched_skills']
                                        score_cv_skills.append(match_skills)

                                    # Add the skills_list as a new column named "skills" to the DataFrame
                                    df['skills'] = skills_list
                                    df['match_score'] = score_cv_scores
                                    df['common_skills'] = score_cv_skills
                                    
                                    # st.session_state.option = 'Choose option'

                                    # Sort the DataFrame by 'match_score' column in descending order
                                    df_sorted = df.sort_values(by='match_score', ascending=False)
                                    threshold_df = df_sorted[df_sorted['match_score'] >= 0.1]
                                    end_time = time.time()
                                    # print("ENNNNNNNNNNNNNNNNNNNNNNNND",end_time)
                                    elapsed_time = end_time - start_time
                                    # st.write(elapsed_time)
                                    if threshold_df.empty:
                                        st.write("No matches above the threshold were found.")
                                    else:
                                        
                                            number_of_resumes = len(df_sorted)
                                            number_of_matches = len(threshold_df)
                                            
                                            fig_pie = visualization.plot_matches_pie(number_of_resumes, number_of_matches)

                                            # Display the pie chart in Streamlit
                                            st.plotly_chart(fig_pie)
                                            st.header(f"Matches above Threshold ")
                                            threshold_df = threshold_df.sort_values(by='match_score', ascending=False)
                          
                                            for index, row in threshold_df.iterrows():
                    
                                                
    
                                                
                                                with st.expander(f"Details for {row['file_name']}"):
                                                        st.markdown("""
                                                        <style>
                                                        .scrollable-container {
                                                            max-height: 400px; /* Adjust the height as needed */
                                                            overflow-y: scroll;
                                                            padding-right: 15px; /* Add padding to avoid cutting off content */
                                                        }
                                                        .scrollable-container::-webkit-scrollbar {
                                                            width: 8px;
                                                        }
                                                        .scrollable-container::-webkit-scrollbar-thumb {
                                                            background-color: #888;
                                                            border-radius: 10px;
                                                        }
                                                        .scrollable-container::-webkit-scrollbar-thumb:hover {
                                                            background-color: #555;
                                                        }
                                                        </style>
                                                        """, unsafe_allow_html=True)

                                                        # Wrap the content in a scrollable container
                                                        # st.markdown('<div class="scrollable-container">', unsafe_allow_html=True)
                                                        # Make the file name clickable to open the PDF file
                                                        file_name = row['file_name']
                                                        upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                                                        # st.markdown(upload_dir_path)
                                                        file_path = os.path.join(upload_dir_path, file_name)     
                                                        # st.markdown(file_path)                                           
                                                        # st.markdown(file_path)
                                                        if os.path.exists(file_path):
                                                            with open(file_path, "rb") as file:
                                                                file_data = file.read()
                                                            st.download_button(label="Open PDF", data=file_data, file_name=file_name, mime="application/pdf")
                                                        else:
                                                            st.write("File not found.")
                                                            
                                                        # URL to access the file
                                                        # file_url = f"uploaded_files/{file_name}"

                                                        # # Display the link
                                                        # st.write(f'<a href="/{file_url}" target="_blank">Open the PDF</a>', unsafe_allow_html=True)
                                                        
                                                        st.write(f"**Match Score:** {row['match_score']}")
                                                        
                                                        # st.markdown('</div>', unsafe_allow_html=True)

                                                        processed_text = row['text']
                                                        essential_info = details(processed_text, row['file_name'])

                                                        st.markdown("""
                                                        <style>
                                                        .dashboard-container {
                                                            padding: 10px;
                                                            background-color: #f0f0f0;
                                                            border-radius: 5px;
                                                            box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                                            margin-top: 20px; /* Add margin to separate from the scrollable container */
                                                        }
                                                        .dashboard-title {
                                                            font-size: 20px;
                                                            font-weight: bold;
                                                            margin-bottom: 10px;
                                                        }
                                                        .dashboard-item {
                                                            margin-bottom: 5px;
                                                        }
                                                        </style>
                                                        """, unsafe_allow_html=True)

                                                        st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
                                                        st.markdown('<div class="dashboard-title">Essential Information</div>', unsafe_allow_html=True)
                                                        for key, value in essential_info.items():
                                                            if key == 'education' and value:
                                                                st.markdown('<div class="dashboard-item"><b>Education:</b></div>', unsafe_allow_html=True)
                                                                education_df = pd.DataFrame(value, columns=['Education'])
                                                                st.dataframe(education_df, width=800)
                                                            elif value:
                                                                st.markdown(f'<div class="dashboard-item"><b>{key}:</b> {value}</div>', unsafe_allow_html=True)
                                                            else:
                                                                st.markdown(f'<div class="dashboard-item"><b>{key}:</b> Not available</div>', unsafe_allow_html=True)
                                                        st.markdown('</div>', unsafe_allow_html=True)
                                                        
                                                        skills_dict = row['skills']
                                                        skills_df = pd.DataFrame.from_dict(skills_dict, orient='index', columns=['type', 'score'])
                                                        # Create a DataFrame with skills from the dictionary
                                                        skills_df = skills_df.sort_values(by='score', ascending=False)
                                                        skills_df_only = pd.DataFrame({'skills': list(skills_dict.keys())})

                                                        # Drop the 'type' column
                                                        skills_df = skills_df.drop(columns=['type'])  # This line drops the 'type' column

                                                        st.subheader("Common Skills with job descritption")
                                                        st.dataframe(skills_df, width=800)
                                                        # Display matched skills DataFrame
                                                        # Ensure matched_skills_df is defined and populated
                                                        if 'common_skills' in row and row['common_skills'] is not None:
                                                            common_skills_list = row['common_skills']
                                                            common_skills_df = pd.DataFrame([common_skills_list], columns=[f"Skill {i+1}" for i in range(len(common_skills_list))])
                                                            # st.subheader("Common Skills with job descritption")
                                                            # st.dataframe(common_skills_df)

                                                        st.markdown("""
                                                        <style>
                                                        .dashboard-container {
                                                            padding: 10px;
                                                            background-color: #ff9999; /* Light red shade */
                                                            border-radius: 5px;
                                                            box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                                            margin-top: 20px; /* Add margin to separate from the scrollable container */
                                                        }
                                                        .dashboard-title {
                                                            font-size: 20px;
                                                            font-weight: bold;
                                                            margin-bottom: 10px;
                                                        }
                                                        .dashboard-item {
                                                            margin-bottom: 5px;
                                                        }
                                                        </style>
                                                        """, unsafe_allow_html=True)

                                                        # Move the pie chart below the essential information and align it to the left
                                                        fig_pie = visualization.display_skills_distribution(skills_df_only, cvs_data)
                                                        st.markdown('<div style="clear:both;"></div>', unsafe_allow_html=True)  # Clear float
                                                        st.plotly_chart(fig_pie, use_container_width=False)
                                                        st.session_state.main_action = None
                                                        st.session_state.option = 'Choose option'
                                                        st.session_state.option1 = 'Choose option'
                                                        # print(st.session_state.option1)
                                                        # print(st.session_state.option)
                                                        
                                            
                                            
                    if main_action == "Details":
                                                st.subheader("Details")
                                                st.markdown("---")

                                                # Convert skills to dictionary
                                                skills_dict = st.session_state.extracted_skills.set_index('full')['score'].to_dict()
                                                
                                                # Display extracted skills and responsibilities side by side in columns
                                                col1, col2 = st.columns(2)

                                                with col1:
                                                    st.subheader("Extracted Skills")
                                                    with st.expander("Skills Details", expanded=True):
                                                        # sorted_skills = st.session_state.skills[['name', 'type', 'score']].sort_values(by='score', ascending=False)

                                                        # Display the sorted DataFrame
                                                        st.dataframe(st.session_state.extracted_skills[['full', 'type', 'score']], height=300)

                                                with col2:
                                                    st.subheader("Extracted Responsibilities")
                                                    with st.expander("Responsibilities Details", expanded=True):
                                                        st.dataframe(st.session_state.extracted_res, height=300)
                                                
                                                if data is None: 
                                                    st.write("data is none")

                                                st.subheader("Skill Categories")
                                                result = visualization.find_skill_in_categories(data, skills_dict)
                                                if result is None: 
                                                    st.write("result is none")
                                                if result:
                                                    categories = []
                                                    subcategories = []
                                                    skills_list = []

                                                    for category, subcategories_data in result.items():
                                                        for subcategory, skills_data in subcategories_data.items():
                                                            if skills_data:  # Check if skills are found in this subcategory
                                                                for skill in skills_data:
                                                                    # print(skill)
                                                                    categories.append(category)
                                                                    subcategories.append(subcategory)
                                                                    skills_list.append(skill['name'])
                                                            else:  # Add placeholders if no skills found in this subcategory
                                                                categories.append(category)
                                                                subcategories.append(subcategory)
                                                                skills_list.append("No skills found")
                                                    with st.expander("Categories", expanded=True):
                                                    #  df = pd.DataFrame({"Category": categories, "Subcategory": subcategories, "Skill": skills_list})
                                                    #  st.dataframe(df, height=300)
                                                    #  visualization.display_hierarchy_tree(categories, subcategories, skills_list)
                                                      visualization.visualize_tree_map(categories, subcategories, skills_list)

                                                st.subheader("Main Subcategories")
                                                with st.expander("View main Subcategories"):
                                    # Add markdown content inside the expander
                                                    # st.markdown("View main categories")
                                                    
                                                    # Visualization using the created function
                                                    visualization.plot_subcategory_skills(result)
                                                    st.session_state.main_action = None
                               
   
            # Display the content based on the sidebar button clicked
    if nav_choice == "Fill Job Description Form":
        
        st.session_state.main_action = None
        # st.header("Fill Job Description Form")
        
        
        job_description_form()
    
    if nav_choice == "Update Resume Database":
        database_file = "pdfs.db"  
        update_database(database_file)
        
        
    if nav_choice == "Additional actions":
        st.sidebar.markdown("---")  # Separator line
        st.sidebar.write("Choose an action:")
        if st.sidebar.button("Update Skills Database"):
            st.empty()
            st.header("Update Skills Database")
            run_script(nlp)

            st.write('Skills Database was updated!')
        if st.sidebar.button("Log Off"):
            st.session_state.logged_in = False
            st.experimental_rerun()

                       
       
    
        
   
def job_description_form():
    
    if 'submitted' not in st.session_state:
        st.session_state.submitted = False
        st.session_state.form_data = None
    
    st.session_state.extracted_skills = None
    st.session_state.extracted_res = None
    
    st.title("Job Description Form")
    with st.form("job_description_form"):
        job_title = st.text_input("Job Title")
        job_description = st.text_area("Job Description")
        responsibilities = st.text_area("Responsibilities", help="Please enter the responsibilities for this job position.")
        skills_required = st.text_area("Skills Required", help="Please enter the skills required for this job position.")
        submit_button = st.form_submit_button("Submit")
    
    if 'main_action2' not in st.session_state:
        st.session_state.main_action2 = None

    if submit_button:
       data = load_tree() 
       if not responsibilities:
            st.error("Responsibilities are required before submitting the form.")
       elif not skills_required:
            st.error("Skills Required are required before submitting the form.")
       else:
        st.session_state.submitted = True
        form_data = {
            "Job Title": job_title,
            "Job Description": job_description,
            "Responsibilities": responsibilities,
            "Skills Required": skills_required
        }

        # Display the submitted data
        
        
        # Perform additional processing
        job_desc = f"{form_data['Job Description']}\n{form_data['Responsibilities']}\n{form_data['Skills Required']}"
        job_desc = translation.translation(job_desc)  # Translate the job description if needed
        # processed_desc = main_test.processing_job(job_desc)  # Process the job description
        st.session_state.extracted_skills1 = main_test.get_skills(job_desc)  # Extract skills
        sentences = preprocessing.tokenize_sentences(preprocessing.special_char(job_desc))
        st.session_state.extracted_res1 = extracting_information.extract_resp(sentences)  # Extract responsibilities

        # Display extracted information in a structured format
        st.success("Job Description Form Submitted Successfully!")
        
    col1, col2, col3 = st.columns(3)

    if col1.button("View Resume Database"):
        st.session_state.main_action2 = "View Resume Database"
    if col2.button("Upload Resumes"):
        st.session_state.main_action2 = "Upload Resumes"
    if col3.button("Details"):
        st.session_state.main_action2 = "Details"

    main_action2 = st.session_state.main_action2

    # Execute main actions based on button clicks
    if main_action2 == "View Resume Database":
        st.session_state.option = 'Choose option'
        database_file = "pdfs.db"  # Replace this with your actual database file path
        view_tables_and_contents(database_file, st.session_state.extracted_skills1)


    if main_action2 == "Upload Resumes":
        st.header("Upload Resumes")

        uploaded_files = st.file_uploader("Choose a file", accept_multiple_files=True, type=["pdf", "docx", "txt", "jpg", "jpeg", "png"])
        
        upload_dir = 'uploaded_files'
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir)
        
        if uploaded_files:
            for uploaded_file in uploaded_files:
                upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                file_path = os.path.join(upload_dir_path, uploaded_file.name)
                with open(file_path, 'wb') as file:
                    file.write(uploaded_file.read())

            df = process_uploaded_pdfs(uploaded_files)
            # updated_df = view_resume_data(df)

            # Display the DataFrame in the Streamlit app
            st.write(df)

        
            processed_data = df["text"]
            skills_list = []
            score_cv_scores = []
            score_cv_skills = []

            # Process each item using list comprehension
            # processed_result = [main_test.processing_job(item) for item in processed_data]

            st.session_state.option = 'Choose option'

            # Display the updated DataFrame
            st.success(f"Processing completed!")

            st.title("Choose Display Option")
            if 'option' not in st.session_state:
                st.session_state.option = 'Choose option'

            # Store the selected option in session state
            selected_option = st.selectbox('Display Options:', ['Choose option', 'Threshold', 'Top Matches'], index=['Choose option', 'Threshold', 'Top Matches'].index(st.session_state.option))
            if selected_option != 'Choose option':
                st.session_state.option = selected_option
                                                
            if st.session_state.option == 'Top Matches':
                
                        # Extract skills from each processed item and add to skills_list
                        cvs_data = dataa[dataa['full'].isin(st.session_state.extracted_skills1['full'])]
                                            # processed_result = [main_test.processing_job(item) for item in processed_data]

                                            # Extract skills from each processed item and add to skills_list
                        for processed_item in processed_data:
                            skills_df = matchers.skills_extraction_pipeline(cvs_data, processed_item, full_matcher, abbv_matcher, nlp)
                            # df_match = skills_df[["full", "type", "score"]]

                            # Create a dictionary for the skills
                            skills_dict = {row["full"]: (row["type"], row["score"]) for _, row in skills_df.iterrows()}
                            skills_list.append(skills_dict)

                            match_score = matching.check_cv_match(st.session_state.extracted_skills1, skills_df)['match_percentage']
                            # print("match score", match_score)
                            score_cv_scores.append(match_score if match_score is not None else 0)
                            # match_skills = matching.check_cv_match(st.session_state.extracted_skills1, skills_df)['matched_skills']
                            # score_cv_skills.append(match_skills)

                        # Add the skills_list as a new column named "skills" to the DataFrame
                        df['skills'] = skills_list
                        df['match_score'] = score_cv_scores
                        # df['common_skills'] = score_cv_skills

                        # Sort the DataFrame by 'match_score' column in descending order
                        df_sorted = df.sort_values(by='match_score', ascending=False)
                        top_5_df = df_sorted.head(5)
                        st.header(f"Top matches ")

                        # Make file names clickable for details
                        for index, row in top_5_df.iterrows():

                            with st.expander(f"Details for {row['file_name']}"):
                                st.markdown("""
                                <style>
                                .scrollable-container {
                                    max-height: 400px; /* Adjust the height as needed */
                                    overflow-y: scroll;
                                    padding-right: 15px; /* Add padding to avoid cutting off content */
                                }
                                .scrollable-container::-webkit-scrollbar {
                                    width: 8px;
                                }
                                .scrollable-container::-webkit-scrollbar-thumb {
                                    background-color: #888;
                                    border-radius: 10px;
                                }
                                .scrollable-container::-webkit-scrollbar-thumb:hover {
                                    background-color: #555;
                                }
                                </style>
                                """, unsafe_allow_html=True)

                                # Wrap the content in a scrollable container
                                st.markdown('<div class="scrollable-container">', unsafe_allow_html=True)
                                file_name = row['file_name']
                                upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                                # st.markdown(upload_dir_path)
                                file_path = os.path.join(upload_dir_path, file_name)     
                                # st.markdown(file_path)                                           
                                # st.markdown(file_path)
                                if os.path.exists(file_path):
                                    with open(file_path, "rb") as file:
                                        file_data = file.read()
                                    st.download_button(label="Open PDF", data=file_data, file_name=file_name, mime="application/pdf")
                                else:
                                    st.write("File not found.")
                                st.write(f"**Match Score:** {row['match_score']}")
                                
                                st.markdown('</div>', unsafe_allow_html=True)

                                processed_text = row['text']
                                essential_info = details(processed_text, row['file_name'])

                                st.markdown("""
                                <style>
                                .dashboard-container {
                                    padding: 10px;
                                    background-color: #f0f0f0;
                                    border-radius: 5px;
                                    box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                    margin-top: 20px; /* Add margin to separate from the scrollable container */
                                }
                                .dashboard-title {
                                    font-size: 20px;
                                    font-weight: bold;
                                    margin-bottom: 10px;
                                }
                                .dashboard-item {
                                    margin-bottom: 5px;
                                }
                                </style>
                                """, unsafe_allow_html=True)

                                st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
                                st.markdown('<div class="dashboard-title">Essential Information</div>', unsafe_allow_html=True)
                                for key, value in essential_info.items():
                                    if key == 'education' and value:
                                        st.markdown('<div class="dashboard-item"><b>Education:</b></div>', unsafe_allow_html=True)
                                        education_df = pd.DataFrame(value, columns=['Education'])
                                        st.dataframe(education_df, width=800)
                                    elif value:
                                        st.markdown(f'<div class="dashboard-item"><b>{key}:</b> {value}</div>', unsafe_allow_html=True)
                                    else:
                                        st.markdown(f'<div class="dashboard-item"><b>{key}:</b> Not available</div>', unsafe_allow_html=True)
                                st.markdown('</div>', unsafe_allow_html=True)
                                
                                skills_dict = row['skills']
                                skills_df = pd.DataFrame.from_dict(skills_dict, orient='index', columns=['type', 'score'])
                                # Create a DataFrame with skills from the dictionary
                                skills_df = skills_df.sort_values(by='score', ascending=False)
                                skills_df_only = pd.DataFrame({'skills': list(skills_dict.keys())})

                                # Drop the 'type' column
                                skills_df = skills_df.drop(columns=['type'])  # This line drops the 'type' column

                                st.subheader("Common Skills with job descritption")
                                st.dataframe(skills_df, width=800)
                                # Display matched skills DataFrame
                                # Ensure matched_skills_df is defined and populated
                                # if 'common_skills' in row and row['common_skills'] is not None:
                                #     common_skills_list = row['common_skills']
                                #     common_skills_df = pd.DataFrame([common_skills_list], columns=[f"Skill {i+1}" for i in range(len(common_skills_list))])
                                #     # st.subheader("Common Skills with job descritption")
                                    # st.dataframe(common_skills_df)

                                st.markdown("""
                                <style>
                                .dashboard-container {
                                    padding: 10px;
                                    background-color: #ff9999; /* Light red shade */
                                    border-radius: 5px;
                                    box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                    margin-top: 20px; /* Add margin to separate from the scrollable container */
                                }
                                .dashboard-title {
                                    font-size: 20px;
                                    font-weight: bold;
                                    margin-bottom: 10px;
                                }
                                .dashboard-item {
                                    margin-bottom: 5px;
                                }
                                </style>
                                """, unsafe_allow_html=True)

                                # Move the pie chart below the essential information and align it to the left
                                fig_pie = visualization.display_skills_distribution(skills_df_only, cvs_data)
                                st.markdown('<div style="clear:both;"></div>', unsafe_allow_html=True)  # Clear float
                                st.plotly_chart(fig_pie, use_container_width=False)
                                st.session_state.main_action2 = None
                                st.session_state.option = 'Choose option'
                                    
            elif st.session_state.option == 'Threshold':
                # threshold_input = st.number_input("Enter the threshold score:", value=0, step=0.01)

                    # Extract skills from each processed item and add to skills_list
                    cvs_data = dataa[dataa['full'].isin(st.session_state.extracted_skills1['full'])]
                                        # processed_result = [main_test.processing_job(item) for item in processed_data]

                                        # Extract skills from each processed item and add to skills_list
                    for processed_item in processed_data:
                        skills_df = matchers.skills_extraction_pipeline(cvs_data, processed_item, full_matcher, abbv_matcher, nlp)
                        # df_match = skills_df[["full", "type", "score"]]

                        # Create a dictionary for the skills
                        skills_dict = {row["full"]: (row["type"], row["score"]) for _, row in skills_df.iterrows()}
                        skills_list.append(skills_dict)

                        match_score = matching.check_cv_match(st.session_state.extracted_skills1, skills_df)['match_percentage']
                        # print("match score", match_score)
                        score_cv_scores.append(match_score if match_score is not None else 0)
                        # match_skills = matching.check_cv_match(st.session_state.extracted_skills1, skills_df)['matched_skills']
                        # score_cv_skills.append(match_skills)

                    # Add the skills_list as a new column named "skills" to the DataFrame
                    df['skills'] = skills_list
                    df['match_score'] = score_cv_scores
                    # df['common_skills'] = score_cv_skills

                    # Sort the DataFrame by 'match_score' column in descending order
                    df_sorted = df.sort_values(by='match_score', ascending=False)
                    threshold_df = df_sorted[df_sorted['match_score'] >= 0.2]
                    if threshold_df.empty:
                        st.write("No matches above the threshold were found.")
                    else:
                        
                            number_of_resumes = len(df_sorted)
                            number_of_matches = len(threshold_df)
                            
                            fig_pie = visualization.plot_matches_pie(number_of_resumes, number_of_matches)

                            # Display the pie chart in Streamlit
                            st.plotly_chart(fig_pie)
                            st.header(f"Matches above Threshold ")    
                            threshold_df = threshold_df.sort_values(by='match_score', ascending=False)                                
                            for index, row in threshold_df.iterrows():

                            
                                
                                with st.expander(f"Details for {row['file_name']}"):
                                        st.markdown("""
                                        <style>
                                        .scrollable-container {
                                            max-height: 400px; /* Adjust the height as needed */
                                            overflow-y: scroll;
                                            padding-right: 15px; /* Add padding to avoid cutting off content */
                                        }
                                        .scrollable-container::-webkit-scrollbar {
                                            width: 8px;
                                        }
                                        .scrollable-container::-webkit-scrollbar-thumb {
                                            background-color: #888;
                                            border-radius: 10px;
                                        }
                                        .scrollable-container::-webkit-scrollbar-thumb:hover {
                                            background-color: #555;
                                        }
                                        </style>
                                        """, unsafe_allow_html=True)

                                        # Wrap the content in a scrollable container
                                        st.markdown('<div class="scrollable-container">', unsafe_allow_html=True)
                                        file_name = row['file_name']
                                        upload_dir_path = os.path.abspath(upload_dir)  # Get the absolute path of upload_dir
                                        # st.markdown(upload_dir_path)
                                        file_path = os.path.join(upload_dir_path, file_name)     
                                        # st.markdown(file_path)                                           
                                        # st.markdown(file_path)
                                        if os.path.exists(file_path):
                                            with open(file_path, "rb") as file:
                                                file_data = file.read()
                                            st.download_button(label="Open PDF", data=file_data, file_name=file_name, mime="application/pdf")
                                        else:
                                            st.write("File not found.")
                                        st.write(f"**Match Score:** {row['match_score']}")
                                        
                                        st.markdown('</div>', unsafe_allow_html=True)

                                        processed_text = row['text']
                                        essential_info = details(processed_text, row['file_name'])

                                        st.markdown("""
                                        <style>
                                        .dashboard-container {
                                            padding: 10px;
                                            background-color: #f0f0f0;
                                            border-radius: 5px;
                                            box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                            margin-top: 20px; /* Add margin to separate from the scrollable container */
                                        }
                                        .dashboard-title {
                                            font-size: 20px;
                                            font-weight: bold;
                                            margin-bottom: 10px;
                                        }
                                        .dashboard-item {
                                            margin-bottom: 5px;
                                        }
                                        </style>
                                        """, unsafe_allow_html=True)

                                        st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
                                        st.markdown('<div class="dashboard-title">Essential Information</div>', unsafe_allow_html=True)
                                        for key, value in essential_info.items():
                                            if key == 'education' and value:
                                                st.markdown('<div class="dashboard-item"><b>Education:</b></div>', unsafe_allow_html=True)
                                                education_df = pd.DataFrame(value, columns=['Education'])
                                                st.dataframe(education_df, width=800)
                                            elif value:
                                                st.markdown(f'<div class="dashboard-item"><b>{key}:</b> {value}</div>', unsafe_allow_html=True)
                                            else:
                                                st.markdown(f'<div class="dashboard-item"><b>{key}:</b> Not available</div>', unsafe_allow_html=True)
                                        st.markdown('</div>', unsafe_allow_html=True)
                                        
                                        skills_dict = row['skills']
                                        skills_df = pd.DataFrame.from_dict(skills_dict, orient='index', columns=['type', 'score'])
                                        # Create a DataFrame with skills from the dictionary
                                        skills_df = skills_df.sort_values(by='score', ascending=False)
                                        skills_df_only = pd.DataFrame({'skills': list(skills_dict.keys())})

                                        # Drop the 'type' column
                                        skills_df = skills_df.drop(columns=['type'])  # This line drops the 'type' column

                                        st.subheader("Common Skills with job descritption")
                                        st.dataframe(skills_df, width=800)
                                        # Display matched skills DataFrame
                                        # Ensure matched_skills_df is defined and populated
                                        # if 'common_skills' in row and row['common_skills'] is not None:
                                        #     common_skills_list = row['common_skills']
                                        #     common_skills_df = pd.DataFrame([common_skills_list], columns=[f"Skill {i+1}" for i in range(len(common_skills_list))])
                                        #     st.subheader("Common Skills with job descritption")
                                        #     st.dataframe(common_skills_df)

                                        st.markdown("""
                                        <style>
                                        .dashboard-container {
                                            padding: 10px;
                                            background-color: #ff9999; /* Light red shade */
                                            border-radius: 5px;
                                            box-shadow: 0px 0px 5px rgba(0, 0, 0, 0.1);
                                            margin-top: 20px; /* Add margin to separate from the scrollable container */
                                        }
                                        .dashboard-title {
                                            font-size: 20px;
                                            font-weight: bold;
                                            margin-bottom: 10px;
                                        }
                                        .dashboard-item {
                                            margin-bottom: 5px;
                                        }
                                        </style>
                                        """, unsafe_allow_html=True)

                                        # Move the pie chart below the essential information and align it to the left
                                        fig_pie = visualization.display_skills_distribution(skills_df_only, cvs_data)
                                        st.markdown('<div style="clear:both;"></div>', unsafe_allow_html=True)  # Clear float
                                        st.plotly_chart(fig_pie, use_container_width=False)
                                        st.session_state.main_action2 = None
                                        st.session_state.option = 'Choose option'
                            
    if main_action2 == "Details":
                                        st.subheader("Details")
                                        st.markdown("---")

                                        # Convert skills to dictionary
                                        skills_dict = st.session_state.extracted_skills1.set_index('full')['score'].to_dict()

                                        # Display extracted skills and responsibilities side by side in columns
                                        col1, col2 = st.columns(2)

                                        with col1:
                                            st.subheader("Extracted Skills")
                                            with st.expander("Skills Details", expanded=True):
                                                st.dataframe(st.session_state.extracted_skills1[['full', 'type', 'score']], height=300)

                                        with col2:
                                            st.subheader("Extracted Responsibilities")
                                            with st.expander("Responsibilities Details", expanded=True):
                                                st.dataframe(st.session_state.extracted_res1, height=300)
                                                
                                        # expander_height = 350  # Adjust this value as needed
                                        # col1.markdown(f"<div style='min-height:{expander_height}px;'></div>", unsafe_allow_html=True)
                                        # col2.markdown(f"<div style='min-height:{expander_height}px;'></div>", unsafe_allow_html=True)

                                        # # Check if data is None
                                        # if data is None: 
                                        #     st.write("Data is none")

                                        # Skill Categories Section
                                        st.subheader("Skill Subcategories")
                                        result = visualization.find_skill_in_categories(data, skills_dict)

                                        if result is None: 
                                            st.write("Result is none")

                                        if result:
                                            categories = []
                                            subcategories = []
                                            skills_list = []

                                            for category, subcategories_data in result.items():
                                                for subcategory, skills_data in subcategories_data.items():
                                                    if skills_data:  # Check if skills are found in this subcategory
                                                        for skill in skills_data:
                                                            categories.append(category)
                                                            subcategories.append(subcategory)
                                                            skills_list.append(skill['name'])
                                                    else:  # Add placeholders if no skills found in this subcategory
                                                        categories.append(category)
                                                        subcategories.append(subcategory)
                                                        skills_list.append("No skills found")

                                            with st.expander("Categories", expanded=True):
                                                # Displaying the tree map for categories, subcategories, and skills
                                                visualization.visualize_tree_map(categories, subcategories, skills_list)

                                        # Main Categories Section
                                        st.subheader("Main Subcategories")
                                        with st.expander("View main Subcategories"):
                                            # Visualization using the created function
                                            visualization.plot_subcategory_skills(result)
                                            st.session_state.main_action2 = None
    

def truncate_text(text, max_chars):
    if len(text) > max_chars:
        return text[:max_chars] + '...'
    return text


# Function to add PDFs to the database
def add_files_to_db(uploaded_file):
    # Connect to the database
    conn = sqlite3.connect('pdfs.db')
    cursor = conn.cursor()
    
    # Define the directory where the uploaded files will be saved
    upload_dir = 'uploaded_DB'
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir)
    
    # Save the uploaded file to the designated directory
    file_path = os.path.join(upload_dir, uploaded_file.name)
    with open(file_path, 'wb') as file:
        file.write(uploaded_file.read())
    
    # Get the real full path of the file
    full_file_path = os.path.abspath(file_path)
    
    # Process the saved file based on its type and extract text
    extracted_text = ""
    if full_file_path.endswith(".pdf"):
        images = convert_from_path(full_file_path)
        for image in images:
            text = pytesseract.image_to_string(image)
            extracted_text += text + "\n"
    elif full_file_path.endswith(".docx"):
        doc = Document(full_file_path)
        extracted_text = "\n".join([para.text for para in doc.paragraphs])
    elif full_file_path.endswith(".txt"):
        with open(full_file_path, 'r') as txt_file:
            extracted_text = txt_file.read()
    elif full_file_path.endswith((".jpg", ".jpeg", ".png")):
        image = Image.open(full_file_path)
        extracted_text = pytesseract.image_to_string(image)
    
    # Insert data into the database
    cursor.execute('''
        INSERT INTO pdfs (file_name, file_path, processed)
        VALUES (?, ?, ?)
    ''', (uploaded_file.name, full_file_path, extracted_text))
    
    # Commit the changes and close the connection
    conn.commit()
    conn.close()
    

# Initialize session state
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False  # Initialize session state if not set

# Check session state for authentication status and display appropriate page
if st.session_state.logged_in:
    show_main_interface()
else:
    login_page()























